from kivy.uix.screenmanager import Screen
from kivy.properties import StringProperty
from datetime import datetime
import os

from db import cur
from context import CURRENT_OWNER
from utils import popup

from docx import Document
from docx.shared import Pt


class EODLedgerScreen(Screen):

    report_text = StringProperty("")
    report_lines = []
    shop_name = ""
    report_date = ""

    def on_pre_enter(self):

        owner_uid = CURRENT_OWNER.get("uid")

        if not owner_uid:
            popup("Session expired", "Please login again")
            self.manager.current = "login_type"
            return

        self.shop_name = self.manager.get_screen("revenue").ids.shop.text
        self.report_date = datetime.now().strftime("%Y-%m-%d")

        self.build_report(self.shop_name, self.report_date)


    def build_report(self, shop, date):

        rows = cur.execute("""
            SELECT particulars, description, amount
            FROM revenue
            WHERE shop_id=? AND date=?
            ORDER BY id
        """, (shop, date)).fetchall()

        income = []
        expense = []

        total_income = 0
        total_expense = 0

        for row in rows:

            particulars, desc, amount = row

            if particulars.lower() == "income":
                income.append((desc, amount))
                total_income += amount

            elif particulars.lower() == "expense":
                expense.append((desc, amount))
                total_expense += amount

        closing = total_income - total_expense

        lines = []

        lines.append(f"SHOP: {shop}")
        lines.append(f"DATE: {date}")
        lines.append("")
        lines.append("INCOME")

        for desc, amt in income:
            lines.append(f"{desc:<25} ₹ {amt:.2f}")

        lines.append(f"Total Income: ₹ {total_income:.2f}")
        lines.append("")
        lines.append("EXPENSE")

        for desc, amt in expense:
            lines.append(f"{desc:<25} ₹ {amt:.2f}")

        lines.append(f"Total Expense: ₹ {total_expense:.2f}")
        lines.append("")
        lines.append(f"CLOSING BALANCE: ₹ {closing:.2f}")

        self.report_lines = lines
        self.report_text = "\n".join(lines)


    def print_report(self):

        try:

            folder = "reports"
            os.makedirs(folder, exist_ok=True)

            filename = f"{folder}/Ledger_{self.shop_name}_{self.report_date}.docx"

            doc = Document()

            header = doc.add_heading("Daily EOD Ledger", level=0)
            header.alignment = 1

            doc.add_paragraph(f"Shop: {self.shop_name}")
            doc.add_paragraph(f"Date: {self.report_date}")

            doc.add_paragraph("")

            for line in self.report_lines:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)

            doc.save(filename)

            popup("Saved", f"Report saved:\n{filename}")

        except Exception as e:

            popup("Error", str(e))


    def close_report(self):

        self.manager.current = "revenue"
